import re
import io
from typing import Optional
import pdfplumber
from docx import Document
from app.config import logger


def normalize_whitespace(text: str) -> str:
    """
    Cleans up excessive whitespace while preserving structure.
    
    Args:
        text: Raw extracted text
    
    Returns:
        Normalized text
    """
    # Replace multiple spaces with single space
    text = re.sub(r' +', ' ', text)
    
    # Replace multiple newlines with double newline (paragraph breaks)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
    
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    return text.strip()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text content from PDF file using pdfplumber.
    
    Args:
        file_bytes: PDF file content as bytes
    
    Returns:
        Extracted text
    
    Raises:
        Exception: If PDF parsing fails
    """
    try:
        text_chunks = []
        
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            if len(pdf.pages) == 0:
                logger.warning("PDF has no pages")
                return ""
            
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                
                if page_text:
                    text_chunks.append(page_text)
                else:
                    logger.warning(f"No text found on page {page_num}")
        
        raw_text = '\n\n'.join(text_chunks)
        return normalize_whitespace(raw_text)
    
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}")
        raise Exception(f"Could not parse PDF file: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from DOCX file using python-docx.
    
    Args:
        file_bytes: DOCX file content as bytes
    
    Returns:
        Extracted text
    
    Raises:
        Exception: If DOCX parsing fails
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        
        if not doc.paragraphs:
            logger.warning("DOCX has no paragraphs")
            return ""
        
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Skip empty paragraphs
                paragraphs.append(text)
        
        # Also extract text from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        
        raw_text = '\n'.join(paragraphs)
        return normalize_whitespace(raw_text)
    
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}")
        raise Exception(f"Could not parse DOCX file: {str(e)}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Routes to appropriate extraction method based on file extension.
    
    Args:
        file_bytes: File content as bytes
        filename: Original filename
    
    Returns:
        Extracted text
    """
    ext = filename.lower().split('.')[-1]
    
    if ext == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif ext == 'docx':
        return extract_text_from_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")
